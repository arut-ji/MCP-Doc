"""Table operation tools"""

from typing import Optional, List
from docx.oxml import OxmlElement

from ..config import logger
from ..processor import processor


def add_table(rows: int, cols: int, data: Optional[List[List[str]]] = None) -> str:
    """
    Add table to document
    
    Parameters:
    - rows: Number of rows
    - cols: Number of columns
    - data: Table data, two-dimensional array
    """
    try:
        if not processor.current_document:
            return "No document is open"
        
        table = processor.current_document.add_table(rows=rows, cols=cols, style="Table Grid")
        
        # Fill table data
        if data:
            for i, row_data in enumerate(data):
                if i < rows:
                    row = table.rows[i]
                    for j, cell_text in enumerate(row_data):
                        if j < cols:
                            row.cells[j].text = str(cell_text)
        
        return f"Added {rows}x{cols} table"
    except Exception as e:
        error_msg = f"Failed to add table: {str(e)}"
        logger.error(error_msg)
        return error_msg


def add_table_row(table_index: int, data: Optional[List[str]] = None) -> str:
    """
    Add a row to table
    
    Parameters:
    - table_index: Table index
    - data: Row data in list format
    """
    try:
        if not processor.current_document:
            return "No document is open"
        
        doc = processor.current_document
        
        if not doc.tables:
            return "No tables in document"
        
        if table_index < 0 or table_index >= len(doc.tables):
            return f"Table index out of range: {table_index}, document has {len(doc.tables)} tables"
        
        table = doc.tables[table_index]
        
        # Add new row
        new_row = table.add_row()
        
        # Fill row data
        if data:
            for i, cell_text in enumerate(data):
                if i < len(new_row.cells):
                    new_row.cells[i].text = str(cell_text)
        
        return f"Added new row to table {table_index}"
    except Exception as e:
        error_msg = f"Failed to add table row: {str(e)}"
        logger.error(error_msg)
        return error_msg


def delete_table_row(table_index: int, row_index: int) -> str:
    """
    Delete a row from table
    
    Parameters:
    - table_index: Table index
    - row_index: Row index to delete
    """
    try:
        if not processor.current_document:
            return "No document is open"
        
        doc = processor.current_document
        
        if not doc.tables:
            return "No tables in document"
        
        if table_index < 0 or table_index >= len(doc.tables):
            return f"Table index out of range: {table_index}, document has {len(doc.tables)} tables"
        
        table = doc.tables[table_index]
        
        if row_index < 0 or row_index >= len(table.rows):
            return f"Row index out of range: {row_index}, table has {len(table.rows)} rows"
        
        # Use XML operations to delete row
        row = table.rows[row_index]._tr
        row.getparent().remove(row)
        
        return f"Deleted row {row_index} from table {table_index}"
    except Exception as e:
        error_msg = f"Failed to delete table row: {str(e)}"
        logger.error(error_msg)
        return error_msg


def edit_table_cell(table_index: int, row_index: int, col_index: int, text: str) -> str:
    """
    Edit table cell content
    
    Parameters:
    - table_index: Table index
    - row_index: Row index
    - col_index: Column index
    - text: Cell text
    """
    try:
        if not processor.current_document:
            return "No document is open"
        
        doc = processor.current_document
        
        if not doc.tables:
            return "No tables in document"
        
        if table_index < 0 or table_index >= len(doc.tables):
            return f"Table index out of range: {table_index}, document has {len(doc.tables)} tables"
        
        table = doc.tables[table_index]
        
        if row_index < 0 or row_index >= len(table.rows):
            return f"Row index out of range: {row_index}, table has {len(table.rows)} rows"
        
        if col_index < 0 or col_index >= len(table.columns):
            return f"Column index out of range: {col_index}, table has {len(table.columns)} columns"
        
        # Modify cell content
        table.cell(row_index, col_index).text = text
        
        return f"Cell ({row_index}, {col_index}) in table {table_index} has been modified"
    except Exception as e:
        error_msg = f"Failed to edit table cell: {str(e)}"
        logger.error(error_msg)
        return error_msg


def merge_table_cells(
    table_index: int,
    start_row: int,
    start_col: int,
    end_row: int,
    end_col: int
) -> str:
    """
    Merge table cells
    
    Parameters:
    - table_index: Table index
    - start_row: Start row index
    - start_col: Start column index
    - end_row: End row index
    - end_col: End column index
    """
    try:
        if not processor.current_document:
            return "No document is open"
        
        doc = processor.current_document
        
        if not doc.tables:
            return "No tables in document"
        
        if table_index < 0 or table_index >= len(doc.tables):
            return f"Table index out of range: {table_index}, document has {len(doc.tables)} tables"
        
        table = doc.tables[table_index]
        
        # Check if row and column indices are valid
        if start_row < 0 or start_row >= len(table.rows):
            return f"Start row index out of range: {start_row}, table has {len(table.rows)} rows"
        
        if start_col < 0 or start_col >= len(table.columns):
            return f"Start column index out of range: {start_col}, table has {len(table.columns)} columns"
        
        if end_row < start_row or end_row >= len(table.rows):
            return f"End row index invalid: {end_row}, should be between {start_row} and {len(table.rows)-1}"
        
        if end_col < start_col or end_col >= len(table.columns):
            return f"End column index invalid: {end_col}, should be between {start_col} and {len(table.columns)-1}"
        
        # Get start and end cells
        start_cell = table.cell(start_row, start_col)
        end_cell = table.cell(end_row, end_col)
        
        # Merge cells
        start_cell.merge(end_cell)
        
        return f"Merged cells in table {table_index} from ({start_row},{start_col}) to ({end_row},{end_col})"
    except Exception as e:
        error_msg = f"Failed to merge table cells: {str(e)}"
        logger.error(error_msg)
        return error_msg


def split_table(table_index: int, row_index: int) -> str:
    """
    Split table into two tables at specified row
    
    Parameters:
    - table_index: Table index
    - row_index: Split table after this row
    """
    try:
        if not processor.current_document:
            return "No document is open"
        
        doc = processor.current_document
        
        if not doc.tables:
            return "No tables in document"
        
        if table_index < 0 or table_index >= len(doc.tables):
            return f"Table index out of range: {table_index}, document has {len(doc.tables)} tables"
        
        table = doc.tables[table_index]
        
        if row_index < 0 or row_index >= len(table.rows) - 1:
            return f"Row index invalid: {row_index}, should be between 0 and {len(table.rows)-2}"
        
        # Use XML operations to split table
        # Get table element
        tbl = table._tbl
        
        # Calculate split position
        split_position = row_index + 1
        
        # Create new table element
        new_tbl = OxmlElement('w:tbl')
        
        # Copy table properties
        for child in tbl.xpath('./w:tblPr')[0].getchildren():
            new_tbl.append(child.copy())
        
        # Copy table grid settings
        for child in tbl.xpath('./w:tblGrid')[0].getchildren():
            new_tbl.append(child.copy())
        
        # Move rows to new table
        rows = tbl.xpath('./w:tr')
        for i in range(split_position, len(rows)):
            new_tbl.append(rows[i])
        
        # Insert new table after original table
        tbl.addnext(new_tbl)
        
        return f"Split table {table_index} after row {row_index}"
    except Exception as e:
        error_msg = f"Failed to split table: {str(e)}"
        logger.error(error_msg)
        return error_msg

