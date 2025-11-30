"""Content editing tools"""

import traceback
from typing import Optional, List, Literal
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

from ..config import logger
from ..processor import processor


def add_paragraph(
    text: str, 
    bold: bool = False, 
    italic: bool = False, 
    underline: bool = False,
    font_size: Optional[int] = None,
    font_name: Optional[str] = None,
    color: Optional[str] = None,
    alignment: Optional[Literal["left", "center", "right", "justify"]] = None
) -> str:
    """
    Add paragraph text to document
    
    Parameters:
    - text: Paragraph text content
    - bold: Whether to bold
    - italic: Whether to italicize
    - underline: Whether to underline
    - font_size: Font size (points)
    - font_name: Font name
    - color: Text color (format: #FF0000)
    - alignment: Alignment (left, center, right, justify)
    """
    try:
        if not processor.current_document:
            return "No document is open"
        
        # Add paragraph
        paragraph = processor.current_document.add_paragraph(text)
        
        # Apply additional formatting
        if paragraph.runs:
            run = paragraph.runs[0]
            run.bold = bold
            run.italic = italic
            run.underline = underline
            
            # Set font size
            if font_size:
                run.font.size = Pt(font_size)
            
            # Set font name
            if font_name:
                run.font.name = font_name
                # Set East Asian font
                run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
            
            # Set font color
            if color and color.startswith('#') and len(color) == 7:
                r = int(color[1:3], 16)
                g = int(color[3:5], 16)
                b = int(color[5:7], 16)
                run.font.color.rgb = RGBColor(r, g, b)
        
        # Set alignment
        if alignment:
            if alignment == "left":
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            elif alignment == "center":
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            elif alignment == "right":
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            elif alignment == "justify":
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        
        return "Paragraph added"
    except Exception as e:
        error_msg = f"Failed to add paragraph: {str(e)}"
        logger.error(error_msg)
        return error_msg


def add_heading(text: str, level: int) -> str:
    """
    Add heading to document
    
    Parameters:
    - text: Heading text
    - level: Heading level (1-9)
    """
    try:
        if not processor.current_document:
            return "No document is open"
        
        processor.current_document.add_heading(text, level=level)
        
        return f"Added level {level} heading"
    except Exception as e:
        error_msg = f"Failed to add heading: {str(e)}"
        logger.error(error_msg)
        return error_msg


def delete_paragraph(paragraph_index: int) -> str:
    """
    Delete specified paragraph from document
    
    Parameters:
    - paragraph_index: Paragraph index to delete
    """
    try:
        if not processor.current_document:
            return "No document is open"
        
        doc = processor.current_document
        
        if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
            return f"Paragraph index out of range: {paragraph_index}, document has {len(doc.paragraphs)} paragraphs"
        
        # python-docx does not provide a direct method to delete a paragraph, use XML operations
        paragraph = doc.paragraphs[paragraph_index]
        p = paragraph._element
        p.getparent().remove(p)
        # Delete paragraph object reference for garbage collection
        paragraph._p = None
        paragraph._element = None
        
        return f"Paragraph {paragraph_index} deleted"
    except Exception as e:
        error_msg = f"Failed to delete paragraph: {str(e)}"
        logger.error(error_msg)
        return error_msg


def delete_text(paragraph_index: int, start_pos: int, end_pos: int) -> str:
    """
    Delete specified text from paragraph
    
    Parameters:
    - paragraph_index: Paragraph index
    - start_pos: Start position (0-based index)
    - end_pos: End position (not included in the text)
    """
    try:
        if not processor.current_document:
            return "No document is open"
        
        doc = processor.current_document
        
        if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
            return f"Paragraph index out of range: {paragraph_index}, document has {len(doc.paragraphs)} paragraphs"
        
        paragraph = doc.paragraphs[paragraph_index]
        text = paragraph.text
        
        if start_pos < 0 or start_pos >= len(text):
            return f"Start position out of range: {start_pos}, paragraph length is {len(text)}"
        
        if end_pos <= start_pos or end_pos > len(text):
            return f"End position invalid: {end_pos}, should be between {start_pos+1} and {len(text)}"
        
        # Build new text (delete specified text)
        new_text = text[:start_pos] + text[end_pos:]
        paragraph.text = new_text
        
        return f"Deleted text from position {start_pos} to {end_pos} in paragraph {paragraph_index}"
    except Exception as e:
        error_msg = f"Failed to delete text: {str(e)}"
        logger.error(error_msg)
        return error_msg


def search_text(keyword: str) -> str:
    """
    Search for text in the document
    
    Parameters:
    - keyword: Keyword to search for
    """
    try:
        if not processor.current_document:
            return "No document is open"
        
        doc = processor.current_document
        results = []
        
        # Search in paragraphs
        for i, paragraph in enumerate(doc.paragraphs):
            if keyword in paragraph.text:
                results.append({
                    "type": "paragraph",
                    "index": i,
                    "text": paragraph.text
                })
        
        # Search in tables
        for t_idx, table in enumerate(doc.tables):
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    if keyword in cell.text:
                        results.append({
                            "type": "table cell",
                            "table_index": t_idx,
                            "row": r_idx,
                            "column": c_idx,
                            "text": cell.text
                        })
        
        if not results:
            return f"Keyword '{keyword}' not found"
        
        # Build response
        response = f"Found {len(results)} occurrences of '{keyword}':\n\n"
        for idx, result in enumerate(results):
            response += f"{idx+1}. {result['type']} "
            if result['type'] == "paragraph":
                response += f"index {result['index']}: {result['text'][:100]}"
                if len(result['text']) > 100:
                    response += "..."
                response += "\n"
            else:
                response += f"in table {result['table_index']} at cell ({result['row']},{result['column']}): {result['text'][:100]}"
                if len(result['text']) > 100:
                    response += "..."
                response += "\n"
        
        return response
    except Exception as e:
        error_msg = f"Failed to search text: {str(e)}"
        logger.error(error_msg)
        return error_msg


def search_and_replace(keyword: str, replace_with: str, preview_only: bool = False) -> str:
    """
    Search and replace text in the document, providing detailed replacement information and preview options
    
    Parameters:
    - keyword: Keyword to search for
    - replace_with: Text to replace with
    - preview_only: Whether to only preview without actually replacing, default is False
    """
    try:
        if not processor.current_document:
            return "No document is open"
        
        doc = processor.current_document
        results = []
        
        # Search in paragraphs
        for i, paragraph in enumerate(doc.paragraphs):
            if keyword in paragraph.text:
                # Save original text and replaced text
                original_text = paragraph.text
                replaced_text = original_text.replace(keyword, replace_with)
                results.append({
                    "type": "paragraph",
                    "index": i,
                    "original": original_text,
                    "replaced": replaced_text,
                    "count": original_text.count(keyword)
                })
                
                # If not in preview mode, perform replacement
                if not preview_only:
                    paragraph.text = replaced_text
        
        # Search in tables
        for t_idx, table in enumerate(doc.tables):
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    if keyword in cell.text:
                        # Save original text and replaced text
                        original_text = cell.text
                        replaced_text = original_text.replace(keyword, replace_with)
                        results.append({
                            "type": "table cell",
                            "table_index": t_idx,
                            "row": r_idx,
                            "column": c_idx,
                            "original": original_text,
                            "replaced": replaced_text,
                            "count": original_text.count(keyword)
                        })
                        
                        # If not in preview mode, perform replacement
                        if not preview_only:
                            # Replace all paragraphs in the cell with the replaced text
                            for para in cell.paragraphs:
                                if keyword in para.text:
                                    para.text = para.text.replace(keyword, replace_with)
        
        if not results:
            return f"Keyword '{keyword}' not found"
        
        # Calculate total replacements
        total_replacements = sum(item["count"] for item in results)
        
        # Build response
        action_word = "Preview" if preview_only else "Replace"
        response = f"{action_word} '{keyword}' with '{replace_with}', found {len(results)} locations, {total_replacements} occurrences:\n\n"
        
        for idx, result in enumerate(results):
            response += f"{idx+1}. In {result['type']} "
            if result['type'] == "paragraph":
                response += f"index {result['index']} {action_word.lower()}ing {result['count']} times:\n"
            else:
                response += f"table {result['table_index']} at cell ({result['row']},{result['column']}) {action_word.lower()}ing {result['count']} times:\n"
            
            # Display original and replaced text snippets (context)
            max_display = 50
            if len(result['original']) > max_display * 2:
                # Find keyword position and display surrounding text
                start_pos = result['original'].find(keyword)
                start_pos = max(0, start_pos - max_display)
                excerpt_original = "..." + result['original'][start_pos:start_pos + max_display * 2] + "..."
                excerpt_replaced = "..." + result['replaced'][start_pos:start_pos + max_display * 2] + "..."
            else:
                excerpt_original = result['original']
                excerpt_replaced = result['replaced']
            
            response += f"  Original: {excerpt_original}\n"
            response += f"  Replaced: {excerpt_replaced}\n\n"
        
        if preview_only:
            response += "This is a preview of replacements. No actual changes were made. To execute replacements, set preview_only to False."
        else:
            response += "Replacements completed successfully."
        
        return response
    except Exception as e:
        error_msg = f"Search and replace failed: {str(e)}"
        logger.error(error_msg)
        return error_msg


def find_and_replace(find_text: str, replace_text: str) -> str:
    """
    Find and replace text in the document
    
    Parameters:
    - find_text: Text to find
    - replace_text: Text to replace with
    """
    try:
        if not processor.current_document:
            return "No document is open"
        
        doc = processor.current_document
        replace_count = 0
        
        # Find and replace in paragraphs
        for paragraph in doc.paragraphs:
            if find_text in paragraph.text:
                paragraph.text = paragraph.text.replace(find_text, replace_text)
                replace_count += paragraph.text.count(replace_text)
        
        # Find and replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if find_text in paragraph.text:
                            paragraph.text = paragraph.text.replace(find_text, replace_text)
                            replace_count += paragraph.text.count(replace_text)
        
        return f"Replaced '{find_text}' with '{replace_text}', {replace_count} occurrences"
    except Exception as e:
        error_msg = f"Find and replace failed: {str(e)}"
        logger.error(error_msg)
        return error_msg


def replace_section(section_title: str, new_content: List[str], preserve_title: bool = True) -> str:
    """
    Find specified title in document and replace content under that title, keeping original position, format, and style
    
    Parameters:
    - section_title: Title text to find
    - new_content: New content list, each element is a paragraph
    - preserve_title: Whether to keep original title, default is True
    """
    try:
        if not processor.current_document:
            return "No document is open"
        
        doc = processor.current_document
        
        # Find title position
        title_index = -1
        for i, paragraph in enumerate(doc.paragraphs):
            if section_title in paragraph.text:
                title_index = i
                break
        
        if title_index == -1:
            return f"Title not found: '{section_title}'"
        
        # Determine end position of that section (next same or higher level title)
        end_index = len(doc.paragraphs)
        title_style = doc.paragraphs[title_index].style
        
        for i in range(title_index + 1, len(doc.paragraphs)):
            # If next same level or higher level title found, set as end position
            if doc.paragraphs[i].style.name.startswith('Heading') and \
               (doc.paragraphs[i].style.name <= title_style.name or doc.paragraphs[i].style == title_style):
                end_index = i
                break
        
        # Save original paragraph style and format information
        original_styles = []
        for i in range(start_delete := (title_index + (1 if preserve_title else 0)), min(end_index, start_delete + len(new_content))):
            if i < len(doc.paragraphs):
                para = doc.paragraphs[i]
                style_info = {
                    'style': para.style,
                    'alignment': para.alignment,
                    'runs': []
                }
                
                # Save each run format
                for run in para.runs:
                    run_info = {
                        'bold': run.bold,
                        'italic': run.italic,
                        'underline': run.underline,
                        'font_size': run.font.size,
                        'font_name': run.font.name,
                        'color': run.font.color.rgb if run.font.color.rgb else None
                    }
                    style_info['runs'].append(run_info)
                
                original_styles.append(style_info)
            else:
                # If original paragraph count is insufficient, use last paragraph style
                if original_styles:
                    original_styles.append(original_styles[-1])
                else:
                    # If no original style, use default style
                    original_styles.append({
                        'style': None,
                        'alignment': None,
                        'runs': []
                    })
        
        # If original style count is insufficient, use last style to fill
        while len(original_styles) < len(new_content):
            if original_styles:
                original_styles.append(original_styles[-1])
            else:
                original_styles.append({
                    'style': None,
                    'alignment': None,
                    'runs': []
                })
        
        # Record insert position
        insert_position = start_delete
        
        # Delete from end to avoid index change
        for i in range(end_index - 1, start_delete - 1, -1):
            p = doc.paragraphs[i]._element
            p.getparent().remove(p)
        
        # Add new content, apply original format
        for i, content in enumerate(reversed(new_content)):
            # Create new paragraph
            p = doc.add_paragraph()
            
            # Apply original paragraph style
            style_info = original_styles[len(new_content) - i - 1]
            if style_info['style']:
                p.style = style_info['style']
            if style_info['alignment'] is not None:
                p.alignment = style_info['alignment']
            
            # Add text and apply format
            if style_info['runs'] and len(style_info['runs']) > 0:
                # If multiple runs, try to keep format
                # Simplified processing: Add entire content to a run, apply format from first run
                run = p.add_run(content)
                run_info = style_info['runs'][0]
                
                run.bold = run_info['bold']
                run.italic = run_info['italic']
                run.underline = run_info['underline']
                
                if run_info['font_size']:
                    run.font.size = run_info['font_size']
                
                if run_info['font_name']:
                    run.font.name = run_info['font_name']
                    # Set Chinese font
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), run_info['font_name'])
                
                if run_info['color']:
                    run.font.color.rgb = run_info['color']
            else:
                # If no run information, add text directly
                p.text = content
            
            # Move new paragraph to correct position
            doc._body._body.insert(insert_position, p._p)
            
            # Delete original added paragraph (at end of document)
            doc._body._body.remove(doc.paragraphs[-1]._p)
        
        return f"Replaced content under title '{section_title}', keeping original format and style"
    except Exception as e:
        error_msg = f"Failed to replace content: {str(e)}"
        logger.error(error_msg)
        traceback.print_exc()  # Print detailed error information
        return error_msg


def edit_section_by_keyword(keyword: str, new_content: List[str], section_range: int = 3) -> str:
    """
    Find paragraphs containing specified keyword and replace them and their surrounding content, keeping original position, format, and style
    
    Parameters:
    - keyword: Keyword to find
    - new_content: New content list, each element is a paragraph
    - section_range: Surrounding paragraph range to replace, default is 3
    """
    try:
        if not processor.current_document:
            return "No document is open"
        
        doc = processor.current_document
        
        # Find keyword position
        keyword_indices = []
        for i, paragraph in enumerate(doc.paragraphs):
            if keyword in paragraph.text:
                keyword_indices.append(i)
        
        if not keyword_indices:
            return f"Keyword not found: '{keyword}'"
        
        # Use first match
        keyword_index = keyword_indices[0]
        
        # Determine paragraph range to replace
        start_index = max(0, keyword_index - section_range)
        end_index = min(len(doc.paragraphs), keyword_index + section_range + 1)
        
        # Save original paragraph style and format information
        original_styles = []
        for i in range(start_index, min(end_index, start_index + len(new_content))):
            if i < len(doc.paragraphs):
                para = doc.paragraphs[i]
                style_info = {
                    'style': para.style,
                    'alignment': para.alignment,
                    'runs': []
                }
                
                # Save each run format
                for run in para.runs:
                    run_info = {
                        'bold': run.bold,
                        'italic': run.italic,
                        'underline': run.underline,
                        'font_size': run.font.size,
                        'font_name': run.font.name,
                        'color': run.font.color.rgb if run.font.color.rgb else None
                    }
                    style_info['runs'].append(run_info)
                
                original_styles.append(style_info)
            else:
                # If original paragraph count is insufficient, use last paragraph style
                if original_styles:
                    original_styles.append(original_styles[-1])
                else:
                    # If no original style, use default style
                    original_styles.append({
                        'style': None,
                        'alignment': None,
                        'runs': []
                    })
        
        # If original style count is insufficient, use last style to fill
        while len(original_styles) < len(new_content):
            if original_styles:
                original_styles.append(original_styles[-1])
            else:
                original_styles.append({
                    'style': None,
                    'alignment': None,
                    'runs': []
                })
        
        # Record insert position
        insert_position = start_index
        
        # Delete from end to avoid index change
        for i in range(end_index - 1, start_index - 1, -1):
            p = doc.paragraphs[i]._element
            p.getparent().remove(p)
        
        # Add new content, apply original format
        for i, content in enumerate(reversed(new_content)):
            # Create new paragraph
            p = doc.add_paragraph()
            
            # Apply original paragraph style
            style_info = original_styles[len(new_content) - i - 1]
            if style_info['style']:
                p.style = style_info['style']
            if style_info['alignment'] is not None:
                p.alignment = style_info['alignment']
            
            # Add text and apply format
            if style_info['runs'] and len(style_info['runs']) > 0:
                # If multiple runs, try to keep format
                # Simplified processing: Add entire content to a run, apply format from first run
                run = p.add_run(content)
                run_info = style_info['runs'][0]
                
                run.bold = run_info['bold']
                run.italic = run_info['italic']
                run.underline = run_info['underline']
                
                if run_info['font_size']:
                    run.font.size = run_info['font_size']
                
                if run_info['font_name']:
                    run.font.name = run_info['font_name']
                    # Set Chinese font
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), run_info['font_name'])
                
                if run_info['color']:
                    run.font.color.rgb = run_info['color']
            else:
                # If no run information, add text directly
                p.text = content
            
            # Move new paragraph to correct position
            doc._body._body.insert(insert_position, p._p)
            
            # Delete original added paragraph (at end of document)
            doc._body._body.remove(doc.paragraphs[-1]._p)
        
        return f"Replaced paragraphs containing keyword '{keyword}' and their surrounding content, keeping original format and style"
    except Exception as e:
        error_msg = f"Failed to replace content: {str(e)}"
        logger.error(error_msg)
        traceback.print_exc()  # Print detailed error information
        return error_msg

