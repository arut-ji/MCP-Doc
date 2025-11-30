"""Document management tools"""

import os
from typing import Optional
from docx import Document
from docx.enum.style import WD_STYLE_TYPE

from ..config import logger
from ..processor import processor


def create_document(file_path: str) -> str:
    """
    Create a new Word document
    
    Parameters:
    - file_path: Document save path
    """
    try:
        processor.current_document = Document()
        processor.current_file_path = file_path
        processor.documents[file_path] = processor.current_document
        
        # Save document
        processor.current_document.save(file_path)
        
        return f"Document created successfully: {file_path}"
    except Exception as e:
        error_msg = f"Failed to create document: {str(e)}"
        logger.error(error_msg)
        return error_msg


def open_document(file_path: str) -> str:
    """
    Open an existing Word document
    
    Parameters:
    - file_path: Path to the document to open
    """
    try:
        if not os.path.exists(file_path):
            return f"File does not exist: {file_path}"
        
        processor.current_document = Document(file_path)
        processor.current_file_path = file_path
        processor.documents[file_path] = processor.current_document
        
        return f"Document opened successfully: {file_path}"
    except Exception as e:
        error_msg = f"Failed to open document: {str(e)}"
        logger.error(error_msg)
        return error_msg


def save_document() -> str:
    """
    Save the currently open Word document to the original file (update the original file)
    """
    try:
        if not processor.current_document:
            return "No document is open"
        
        if not processor.current_file_path:
            return "Current document has not been saved before, please use save_as_document to specify a save path"
            
        # Save to original file path
        processor.current_document.save(processor.current_file_path)
        
        return f"Document saved successfully to original file: {processor.current_file_path}"
    except Exception as e:
        error_msg = f"Failed to save document: {str(e)}"
        logger.error(error_msg)
        return error_msg


def save_as_document(new_file_path: str) -> str:
    """
    Save current document as a new file
    
    Parameters:
    - new_file_path: Path to save the new file
    """
    try:
        if not processor.current_document:
            return "No document is open"
        
        # Save as new file
        processor.current_document.save(new_file_path)
        
        # Update current file path
        processor.current_file_path = new_file_path
        processor.documents[new_file_path] = processor.current_document
        
        return f"Document saved as: {new_file_path}"
    except Exception as e:
        error_msg = f"Failed to save document: {str(e)}"
        logger.error(error_msg)
        return error_msg


def create_document_copy(suffix: str = "-副本") -> str:
    """
    Create a copy of the current document in the directory of the original file
    
    Parameters:
    - suffix: Suffix to add to the original file name, default is "-副本"
    """
    try:
        if not processor.current_document:
            return "No document is open"
        
        if not processor.current_file_path:
            return "Current document has not been saved, cannot create a copy"
        
        # Parse original file path
        file_dir = os.path.dirname(processor.current_file_path)
        file_name = os.path.basename(processor.current_file_path)
        file_name_without_ext, file_ext = os.path.splitext(file_name)
        
        # Create new file name
        new_file_name = f"{file_name_without_ext}{suffix}{file_ext}"
        new_file_path = os.path.join(file_dir, new_file_name)
        
        # Save as new file
        processor.current_document.save(new_file_path)
        
        return f"Document copy created: {new_file_path}"
    except Exception as e:
        error_msg = f"Failed to create document copy: {str(e)}"
        logger.error(error_msg)
        return error_msg


def get_document_info() -> str:
    """
    Get document information, including paragraph count, table count, styles, etc.
    """
    try:
        if not processor.current_document:
            return "No document is open"
        
        doc = processor.current_document
        
        # Get basic document information
        sections_count = len(doc.sections)
        paragraphs_count = len(doc.paragraphs)
        tables_count = len(doc.tables)
        
        # Get style list
        paragraph_styles = []
        for style in doc.styles:
            if style.type == WD_STYLE_TYPE.PARAGRAPH:
                paragraph_styles.append(style.name)
        
        # Build information string
        info = f"Document path: {processor.current_file_path}\n"
        info += f"Section count: {sections_count}\n"
        info += f"Paragraph count: {paragraphs_count}\n"
        info += f"Table count: {tables_count}\n"
        info += f"Available paragraph styles: {', '.join(paragraph_styles[:10])}..."
        
        return info
    except Exception as e:
        error_msg = f"Failed to get document information: {str(e)}"
        logger.error(error_msg)
        return error_msg

